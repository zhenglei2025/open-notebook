"""
Markdown-to-DOCX conversion utility.

Converts markdown text into a formatted Word (.docx) document
with proper headings, bold, italic, lists, tables, code blocks, etc.
"""

import io
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _add_formatted_run(paragraph, text: str, bold: bool = False, italic: bool = False,
                       font_name: str | None = None, font_size: Pt | None = None,
                       color: RGBColor | None = None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_name:
        run.font.name = font_name
        # Also set East Asian font for CJK support
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = r.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size
    if color:
        run.font.color.rgb = color
    return run


def _parse_inline(paragraph, text: str, bold: bool = False, italic: bool = False):
    """Parse inline markdown formatting (bold, italic, code, links) and add runs."""
    # Pattern order matters: bold+italic first, then bold, italic, inline code, links
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'    # ***bold italic***
        r'|(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
        r'|(`([^`]+)`)'           # `code`
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # [text](url)
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > last_end:
            _add_formatted_run(paragraph, text[last_end:m.start()], bold=bold, italic=italic)

        if m.group(2):  # bold+italic
            _add_formatted_run(paragraph, m.group(2), bold=True, italic=True)
        elif m.group(4):  # bold
            _add_formatted_run(paragraph, m.group(4), bold=True, italic=italic)
        elif m.group(6):  # italic
            _add_formatted_run(paragraph, m.group(6), bold=bold, italic=True)
        elif m.group(8):  # inline code
            _add_formatted_run(paragraph, m.group(8), font_name='Courier New',
                             font_size=Pt(9), color=RGBColor(0x60, 0x60, 0x60))
        elif m.group(10):  # link
            link_text = m.group(10)
            _add_formatted_run(paragraph, link_text, color=RGBColor(0x09, 0x69, 0xDA))

        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        _add_formatted_run(paragraph, text[last_end:], bold=bold, italic=italic)


def _parse_table(doc: Document, lines: list[str], start_idx: int) -> int:
    """Parse a markdown table starting at start_idx, return the index after the table."""
    table_lines = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        # Skip separator rows (---|----|---)
        if re.match(r'^[\s\-:|]+$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')]
        table_lines.append(cells)
        i += 1

    if not table_lines:
        return start_idx

    # Determine column count
    max_cols = max(len(row) for row in table_lines)

    table = doc.add_table(rows=len(table_lines), cols=max_cols)
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(table_lines):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < max_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                _parse_inline(p, cell_text, bold=(row_idx == 0))

    return i


def markdown_to_docx(markdown_text: str, title: Optional[str] = None) -> io.BytesIO:
    """
    Convert markdown text to a formatted .docx document.

    Args:
        markdown_text: The markdown content to convert.
        title: Optional title for the document.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    # Set East Asian font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Override heading styles to use black color instead of default blue
    for heading_level in range(1, 5):
        heading_style = doc.styles[f'Heading {heading_level}']
        heading_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading_style.font.name = 'Microsoft YaHei'
        h_rPr = heading_style.element.get_or_add_rPr()
        h_rFonts = h_rPr.find(qn('w:rFonts'))
        if h_rFonts is None:
            h_rFonts = heading_style.element.makeelement(qn('w:rFonts'), {})
            h_rPr.insert(0, h_rFonts)
        h_rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Add title if provided
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Clean citation markers [source:xxx], [note:xxx], [insight:xxx]
    markdown_text = re.sub(r'\[(source|note|insight):[^\]]+\]', '', markdown_text)

    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_block_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                _add_formatted_run(p, code_text, font_name='Courier New', font_size=Pt(9),
                                 color=RGBColor(0x33, 0x33, 0x33))
                # Add background shading
                pPr = p._element.get_or_add_pPr()
                shd = p._element.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F6F8FA',
                })
                pPr.append(shd)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = p._element.makeelement(qn('w:pBdr'), {})
            bottom = p._element.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'D0D7DE',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table detection
        if '|' in stripped and i + 1 < len(lines) and '|' in lines[i + 1]:
            i = _parse_table(doc, lines, i)
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _parse_inline(p, quote_text, italic=True)
            # Color the runs gray
            for run in p.runs:
                if not run.font.color.rgb:
                    run.font.color.rgb = RGBColor(0x65, 0x6D, 0x76)
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^(\s*)[*\-+]\s+(.+)$', stripped)
        if ul_match:
            list_text = ul_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            _parse_inline(p, list_text)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', stripped)
        if ol_match:
            list_text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            _parse_inline(p, list_text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _parse_inline(p, stripped)
        i += 1

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
